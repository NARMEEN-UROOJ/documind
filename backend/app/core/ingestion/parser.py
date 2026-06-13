import fitz
import io
import logging
from pathlib import Path
from typing import Union

logger = logging.getLogger(__name__)


def parse_file(
    file_input: Union[bytes, str, Path],
    source_name: str = "document"
) -> list[dict]:
    ext = Path(source_name).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_input, source_name)
    elif ext in (".txt", ".md"):
        return _parse_text(file_input, source_name)
    elif ext == ".docx":
        return _parse_docx(file_input, source_name)
    else:
        raise ValueError(f"Unsupported format '{ext}'. Use: PDF, TXT, MD, DOCX")


def _parse_pdf(file_input, source_name):
    try:
        doc = fitz.open(stream=file_input, filetype="pdf") \
            if isinstance(file_input, bytes) \
            else fitz.open(str(file_input))
    except Exception as e:
        raise ValueError(f"Could not parse PDF: {e}")

    total = len(doc)
    pages = [
        {"page_num": i + 1, "text": doc[i].get_text().strip(),
         "source": source_name, "total_pages": total}
        for i in range(total) if doc[i].get_text().strip()
    ]
    doc.close()
    logger.info(f"PDF '{source_name}': {len(pages)}/{total} pages")
    return pages


def _parse_text(file_input, source_name):
    text = file_input.decode("utf-8", errors="ignore") \
        if isinstance(file_input, bytes) \
        else Path(file_input).read_text(encoding="utf-8", errors="ignore")

    CHUNK = 3000
    lines, pages_raw, current, length = text.splitlines(), [], [], 0
    for line in lines:
        current.append(line); length += len(line)
        if length >= CHUNK:
            pages_raw.append("\n".join(current).strip())
            current, length = [], 0
    if current:
        pages_raw.append("\n".join(current).strip())

    pages = [{"page_num": i + 1, "text": t, "source": source_name,
               "total_pages": len(pages_raw)}
             for i, t in enumerate(pages_raw) if t]
    logger.info(f"Text '{source_name}': {len(pages)} sections")
    return pages


def _parse_docx(file_input, source_name):
    from docx import Document
    doc = Document(io.BytesIO(file_input)) \
        if isinstance(file_input, bytes) else Document(str(file_input))

    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    groups = [paras[i:i + 50] for i in range(0, len(paras), 50)]
    pages = [{"page_num": i + 1, "text": "\n".join(g),
               "source": source_name, "total_pages": len(groups)}
             for i, g in enumerate(groups) if g]
    logger.info(f"DOCX '{source_name}': {len(pages)} page groups")
    return pages